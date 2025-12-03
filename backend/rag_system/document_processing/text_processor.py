"""
Text Processor
=============

Procesador de documentos de texto (DOCX, TXT, MD) que preserva
el contexto y estructura para el sistema RAG.
"""

import os
import logging
from typing import List, Dict, Any, Optional, Tuple
import re
from pathlib import Path
from datetime import datetime

try:
    from docx import Document
    PYTHON_DOCX_AVAILABLE = True
except ImportError:
    PYTHON_DOCX_AVAILABLE = False

try:
    import markdown
    MARKDOWN_AVAILABLE = True
except ImportError:
    MARKDOWN_AVAILABLE = False

from ..config import DOCUMENT_PROCESSING_CONFIG

logger = logging.getLogger(__name__)

class TextProcessor:
    """
    Procesador de documentos de texto en múltiples formatos
    """
    
    def __init__(self):
        self.config = DOCUMENT_PROCESSING_CONFIG
        self.chunk_size = self.config['chunk_size']
        self.chunk_overlap = self.config['chunk_overlap']
        self.preserve_structure = self.config['preserve_structure']
        self.extract_metadata = self.config['extract_metadata']
        
        self.supported_formats = {
            '.txt': self._process_txt,
            '.md': self._process_markdown,
            '.docx': self._process_docx,
            '.doc': self._process_docx,  # Intentará procesar como DOCX
        }
    
    def is_available(self) -> bool:
        """Verificar disponibilidad básica (TXT siempre disponible)"""
        return True
    
    def get_supported_formats(self) -> List[str]:
        """Obtener lista de formatos soportados"""
        formats = ['.txt', '.md']
        if PYTHON_DOCX_AVAILABLE:
            formats.extend(['.docx', '.doc'])
        return formats
    
    def _get_file_metadata(self, file_path: str) -> Dict[str, Any]:
        """
        Extraer metadatos básicos del archivo
        
        Args:
            file_path: Ruta al archivo
        
        Returns:
            Diccionario con metadatos
        """
        metadata = {}
        
        if self.extract_metadata and os.path.exists(file_path):
            stat = os.stat(file_path)
            metadata = {
                'file_name': os.path.basename(file_path),
                'file_path': file_path,
                'file_size': stat.st_size,
                'creation_date': datetime.fromtimestamp(stat.st_ctime).isoformat(),
                'modification_date': datetime.fromtimestamp(stat.st_mtime).isoformat(),
                'file_extension': os.path.splitext(file_path)[1].lower()
            }
        
        return metadata
    
    def _clean_text(self, text: str) -> str:
        """
        Limpiar y normalizar texto
        
        Args:
            text: Texto a limpiar
        
        Returns:
            Texto limpio
        """
        if not text:
            return ""
        
        # Normalizar espacios en blanco (múltiples espacios -> uno solo)
        text = re.sub(r'[ \t]+', ' ', text)
        
        # Eliminar caracteres de control excepto saltos de línea y tabs
        text = re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F\x7F]', '', text)
        
        # Normalizar saltos de línea múltiples (más de 2 seguidos -> 2)
        text = re.sub(r'\n\s*\n\s*\n+', '\n\n', text)
        
        # Eliminar espacios al inicio y final de líneas
        lines = [line.strip() for line in text.split('\n')]
        text = '\n'.join(lines)
        
        return text.strip()
        
        return text.strip()
    
    def _process_txt(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """
        Procesar archivo TXT
        
        Args:
            file_path: Ruta al archivo TXT
        
        Returns:
            Tupla con (texto, metadatos)
        """
        metadata = self._get_file_metadata(file_path)
        
        try:
            # Intentar diferentes codificaciones
            encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        text = file.read()
                    
                    metadata['encoding'] = encoding
                    break
                except UnicodeDecodeError:
                    continue
            else:
                raise UnicodeDecodeError("No se pudo decodificar el archivo con ninguna codificación")
            
            # Limpiar texto
            text = self._clean_text(text)
            
            # Metadatos adicionales
            if self.extract_metadata:
                lines = text.split('\n')
                metadata.update({
                    'line_count': len(lines),
                    'word_count': len(text.split()),
                    'char_count': len(text)
                })
            
            logger.info(f"✅ Archivo TXT procesado: {len(text)} caracteres")
            
            return text, metadata
            
        except Exception as e:
            logger.error(f"❌ Error procesando TXT {file_path}: {e}")
            raise
    
    def _process_markdown(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """
        Procesar archivo Markdown
        
        Args:
            file_path: Ruta al archivo MD
        
        Returns:
            Tupla con (texto, metadatos)
        """
        metadata = self._get_file_metadata(file_path)
        
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                markdown_text = file.read()
            
            # Si tenemos markdown disponible, convertir a texto plano
            if MARKDOWN_AVAILABLE and not self.preserve_structure:
                # Convertir markdown a HTML y luego extraer texto
                html = markdown.markdown(markdown_text)
                # Remover tags HTML básicos
                text = re.sub(r'<[^>]+>', '', html)
            else:
                # Mantener estructura markdown o procesarlo como texto plano
                text = markdown_text
                
                if self.preserve_structure:
                    # Agregar marcadores de estructura
                    text = self._enhance_markdown_structure(text)
            
            text = self._clean_text(text)
            
            # Metadatos específicos de Markdown
            if self.extract_metadata:
                headers = re.findall(r'^#+\\s+(.+)$', markdown_text, re.MULTILINE)
                links = re.findall(r'\\[([^\\]]+)\\]\\([^\\)]+\\)', markdown_text)
                
                metadata.update({
                    'format': 'markdown',
                    'header_count': len(headers),
                    'headers': headers[:10],  # Primeros 10 headers
                    'link_count': len(links),
                    'word_count': len(text.split()),
                    'char_count': len(text)
                })
            
            logger.info(f"✅ Archivo Markdown procesado: {len(text)} caracteres")
            
            return text, metadata
            
        except Exception as e:
            logger.error(f"❌ Error procesando Markdown {file_path}: {e}")
            raise
    
    def _enhance_markdown_structure(self, text: str) -> str:
        """
        Mejorar estructura de Markdown para preservar contexto
        
        Args:
            text: Texto markdown
        
        Returns:
            Texto con estructura mejorada
        """
        # Marcar headers con contexto
        text = re.sub(r'^(#{1,6})\s+(.+)$', r'\n\n=== HEADER \1 ===\n\2\n', text, flags=re.MULTILINE)
        
        # Marcar listas
        text = re.sub(r'^(\s*[-*+])\s+(.+)$', r'• \2', text, flags=re.MULTILINE)
        text = re.sub(r'^(\s*\d+\.)\s+(.+)$', r'\1 \2', text, flags=re.MULTILINE)
        
        # Marcar código
        text = re.sub(r'```([^\n]*)\n([^`]+)```', r'\n\n=== CÓDIGO \1 ===\n\2\n=== FIN CÓDIGO ===\n', text)
        text = re.sub(r'`([^`]+)`', r'<CÓDIGO: \1>', text)
        
        return text
    
    def _process_docx(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """
        Procesar archivo DOCX
        
        Args:
            file_path: Ruta al archivo DOCX
        
        Returns:
            Tupla con (texto, metadatos)
        """
        if not PYTHON_DOCX_AVAILABLE:
            raise ImportError("python-docx no está disponible para procesar archivos DOCX")
        
        metadata = self._get_file_metadata(file_path)
        
        try:
            doc = Document(file_path)
            
            text_parts = []
            
            # Extraer metadatos del documento
            if self.extract_metadata:
                core_props = doc.core_properties
                metadata.update({
                    'format': 'docx',
                    'title': core_props.title or '',
                    'author': core_props.author or '',
                    'subject': core_props.subject or '',
                    'created': core_props.created.isoformat() if core_props.created else '',
                    'modified': core_props.modified.isoformat() if core_props.modified else '',
                    'category': core_props.category or '',
                    'comments': core_props.comments or '',
                    'keywords': core_props.keywords or ''
                })
            
            # Extraer texto de párrafos
            for para in doc.paragraphs:
                if para.text.strip():
                    if self.preserve_structure and para.style.name.startswith('Heading'):
                        # Marcar headers
                        level = para.style.name.replace('Heading', '').strip()
                        text_parts.append(f"\n\n=== HEADER {level} ===\n{para.text}\n")
                    else:
                        text_parts.append(para.text)
            
            # Extraer texto de tablas
            for table in doc.tables:
                if self.preserve_structure:
                    text_parts.append("\n\n=== TABLA ===\n")
                
                for row in table.rows:
                    row_texts = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_texts.append(cell.text.strip())
                    
                    if row_texts:
                        text_parts.append(" | ".join(row_texts))
                
                if self.preserve_structure:
                    text_parts.append("=== FIN TABLA ===\n")
            
            text = "\n".join(text_parts)
            text = self._clean_text(text)
            
            # Metadatos adicionales
            if self.extract_metadata:
                metadata.update({
                    'paragraph_count': len(doc.paragraphs),
                    'table_count': len(doc.tables),
                    'word_count': len(text.split()),
                    'char_count': len(text)
                })
            
            logger.info(f"✅ Archivo DOCX procesado: {len(text)} caracteres")
            
            return text, metadata
            
        except Exception as e:
            logger.error(f"❌ Error procesando DOCX {file_path}: {e}")
            raise
    
    def extract_text(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """
        Extraer texto del archivo según su formato
        
        Args:
            file_path: Ruta al archivo
        
        Returns:
            Tupla con (texto_extraído, metadatos)
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"Archivo no encontrado: {file_path}")
        
        file_ext = os.path.splitext(file_path)[1].lower()
        
        if file_ext not in self.supported_formats:
            raise ValueError(f"Formato no soportado: {file_ext}")
        
        logger.info(f"📄 Procesando archivo {file_ext}: {os.path.basename(file_path)}")
        
        processor_func = self.supported_formats[file_ext]
        return processor_func(file_path)
    
    def chunk_text(self, text: str, metadata: Dict[str, Any]) -> List[Dict[str, Any]]:
        """
        Dividir texto en chunks para procesamiento RAG
        
        Args:
            text: Texto completo del documento
            metadata: Metadatos del documento
        
        Returns:
            Lista de chunks con metadatos
        """
        if not text.strip():
            return []
        
        chunks = []
        
        # Dividir por secciones si se preserva estructura
        if self.preserve_structure:
            # Buscar divisiones naturales (headers, tablas, etc.)
            sections = re.split(r'\\n\\n(?:===|---)', text)
            
            for section_num, section in enumerate(sections):
                section = section.strip()
                if not section:
                    continue
                
                # Dividir sección en chunks si es muy larga
                section_chunks = self._split_text_into_chunks(section)
                
                for chunk_num, chunk_text in enumerate(section_chunks):
                    chunk = {
                        'text': chunk_text,
                        'metadata': {
                            **metadata,
                            'chunk_id': len(chunks),
                            'section_num': section_num,
                            'chunk_num': chunk_num,
                            'total_chunks': len(section_chunks),
                            'char_count': len(chunk_text),
                            'word_count': len(chunk_text.split())
                        }
                    }
                    chunks.append(chunk)
        else:
            # División simple por tamaño
            text_chunks = self._split_text_into_chunks(text)
            
            for chunk_num, chunk_text in enumerate(text_chunks):
                chunk = {
                    'text': chunk_text,
                    'metadata': {
                        **metadata,
                        'chunk_id': chunk_num,
                        'total_chunks': len(text_chunks),
                        'char_count': len(chunk_text),
                        'word_count': len(chunk_text.split())
                    }
                }
                chunks.append(chunk)
        
        logger.info(f"📝 Texto dividido en {len(chunks)} chunks")
        return chunks
    
    def _split_text_into_chunks(self, text: str) -> List[str]:
        """
        Dividir texto en chunks respetando límites de palabras
        
        Args:
            text: Texto a dividir
        
        Returns:
            Lista de chunks de texto
        """
        if len(text) <= self.chunk_size:
            return [text]
        
        chunks = []
        start = 0
        
        while start < len(text):
            end = start + self.chunk_size
            
            if end >= len(text):
                # Último chunk
                chunks.append(text[start:])
                break
            
            # Buscar un buen punto de corte (espacio o puntuación)
            cut_point = end
            for i in range(end, max(start + self.chunk_size - 200, start), -1):
                if text[i] in ' \\n.!?;':
                    cut_point = i + 1
                    break
            
            chunks.append(text[start:cut_point])
            start = cut_point - self.chunk_overlap
        
        return chunks
    
    def process_document(self, file_path: str) -> Dict[str, Any]:
        """
        Procesar completamente un documento de texto
        
        Args:
            file_path: Ruta al archivo
        
        Returns:
            Resultado completo del procesamiento
        """
        try:
            # Extraer texto y metadatos
            text, metadata = self.extract_text(file_path)
            
            if not text.strip():
                return {
                    "success": False,
                    "error": "No se pudo extraer texto del documento",
                    "file_path": file_path
                }
            
            # Dividir en chunks
            chunks = self.chunk_text(text, metadata)
            
            result = {
                "success": True,
                "file_path": file_path,
                "file_name": os.path.basename(file_path),
                "text": text,
                "metadata": metadata,
                "chunks": chunks,
                "stats": {
                    "total_chars": len(text),
                    "total_words": len(text.split()),
                    "total_chunks": len(chunks),
                    "format": metadata.get('format', os.path.splitext(file_path)[1])
                }
            }
            
            logger.info(f"✅ Documento procesado exitosamente: {result['stats']}")
            return result
            
        except Exception as e:
            logger.error(f"❌ Error procesando documento {file_path}: {e}")
            return {
                "success": False,
                "error": str(e),
                "file_path": file_path
            }
    
    def validate_file(self, file_path: str) -> bool:
        """
        Validar que el archivo sea procesable
        
        Args:
            file_path: Ruta al archivo
        
        Returns:
            True si es válido
        """
        if not os.path.exists(file_path):
            return False
        
        file_ext = os.path.splitext(file_path)[1].lower()
        
        if file_ext not in self.supported_formats:
            return False
        
        # Verificar que no esté vacío
        try:
            return os.path.getsize(file_path) > 0
        except OSError:
            return False